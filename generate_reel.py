from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
GOLD    = RGBColor(0xC9, 0xA0, 0x2C)
NAVY    = RGBColor(0x1B, 0x2A, 0x4A)
CHARCOAL = RGBColor(0x36, 0x36, 0x36)
BLACK   = RGBColor(0x00, 0x00, 0x00)

def set_rtl(paragraph):
    """Make a paragraph right-to-left."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    run.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(6)
    return p

def add_section_heading(doc, text, colour=GOLD):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f'▌ {text}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = colour
    run.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_label(doc, label_en, align_right=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align_right else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(label_en)
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = CHARCOAL
    run.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_arabic(doc, text, timing=None):
    """Add right-aligned Arabic paragraph at 14pt."""
    if timing:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        tr = tp.add_run(f'[{timing}]')
        tr.italic = True
        tr.font.size = Pt(9)
        tr.font.color.rgb = GOLD
        tr.font.name = 'Arial'
        set_rtl(tp)
        tp.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = BLACK
    run.font.name = 'Arial'
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    return p

def add_english(doc, text, bold=False, colour=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(12)
    run.font.color.rgb = colour or BLACK
    run.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(4)
    return p

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C9A02C')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT STARTS
# ══════════════════════════════════════════════════════════════════════════════

add_title(doc, 'Atlas for Humanity — Daily Reel')
add_title(doc, '16 June 2026 | The Queue')
add_hr(doc)

# ── STEP 1: Research Summary ───────────────────────────────────────────────────
add_section_heading(doc, 'STEP 1 — RESEARCH SUMMARY', NAVY)

add_english(doc, 'Trending UK Social Media (last 48 hrs):', bold=True)
add_english(doc, (
    '• British culture and identity content is surging on TikTok UK and Instagram Reels, with '
    '#GrowingUpBritish, #BritishHumour and #BritishMemes consistently generating high engagement.\n'
    '• Summer 2026 "scorching start" social wave: local, community-driven content is outperforming '
    'polished brand content.\n'
    '• Reaction / "decode" formats (creator reacts to a relatable situation) are the dominant '
    'Reels structure right now — perfectly matched to Atlas\'s format.\n'
    '• The MJ biopic and Charli XCX Brat aesthetic are the dominant audio trends, but the '
    '"third place / home" and confessional POV formats suit cultural storytelling.\n'
    'Sources: socialbee.com/blog/instagram-trends | startups.co.uk/news/social-media-trends-june-2026'
))

add_english(doc, 'UK Media / Cultural Moment:', bold=True)
add_english(doc, (
    '• BBC announced 10 % cost cuts — media and public institutions under pressure; '
    'conversations about "British identity" and what Britain stands for are spiking.\n'
    '• Summer heatwave arriving: Brits queuing for ice cream, outdoor events, festivals. '
    'Queuing content is hyper-relevant right now.\n'
    'Source: BBC / ONS public trends June 2026'
))

add_english(doc, 'Strongest Hook Identified:', bold=True)
add_english(doc, (
    'The Queue. The single most recognisable British social institution — and the single biggest '
    'culture-shock moment for almost every SWANA person who arrives in the UK. '
    'Timing: summer heatwave + outdoor queuing season makes this land perfectly.'
))

add_english(doc, 'Real-Life Story Anchor:', bold=True)
add_english(doc, (
    'WWII London rationing queues, 1940–1945. The British government deliberately turned '
    '"queuing for your rations" into a moral and patriotic act. Propaganda posters said '
    '"queue and be fair." Women queued for hours outside butchers and fishmongers. '
    'Jumping the queue was not just rude — it was considered unpatriotic, almost treasonous. '
    'This is the moment the queue became sacred in British culture.\n'
    'Photo: Wikimedia Commons — Category: Rationing in the United Kingdom in World War II\n'
    'URL: https://commons.wikimedia.org/wiki/Category:Rationing_in_the_United_Kingdom_in_World_War_II\n'
    'Source: Debrett\'s / BBC Magazine / Parsons Green Prep cultural history pieces'
))

add_english(doc, 'SWANA / Egyptian Bridge:', bold=True)
add_english(doc, (
    'الواسطة (wasta) — using connections, relationships, and social capital to navigate systems. '
    'In Egypt, this is not cheating; it is community and warmth. The contrast with British queuing '
    'is not "wrong vs right" — it is two completely different philosophies of how society should '
    'allocate limited resources. This frame avoids cliché and creates a genuine "discovery" moment.'
))

add_hr(doc)

# ── STEP 2: Content Architecture ──────────────────────────────────────────────
add_section_heading(doc, 'STEP 2 — CONTENT ARCHITECTURE', NAVY)

rows = [
    ('Topic chosen', 'The British Queue — "queue jumping" as social crime'),
    ('Why today', 'Summer heatwave + outdoor season; British identity content peaking on UK TikTok/IG; decode-format Reels dominating'),
    ('Core cultural insight', 'The queue is not about politeness — it is about WWII-era rationing morality baked into British identity. Understanding this changes how you read every British social situation.'),
    ('Real-life story anchor', 'London WWII rationing queues, 1940–45. Government propaganda made queuing a patriotic duty. Source: commons.wikimedia.org/wiki/Category:Rationing_in_the_United_Kingdom_in_World_War_II'),
    ('SWANA cultural bridge', 'الواسطة — wasta — the Egyptian / SWANA practice of navigating systems through relationships. Not laziness or corruption: community intelligence. Two systems, two philosophies.'),
    ('Photograph / visual', 'https://commons.wikimedia.org/wiki/Category:Rationing_in_the_United_Kingdom_in_World_War_II (women queuing outside London fishmonger, 1945)'),
    ('Reel duration estimate', '75–90 seconds'),
]
for label, val in rows:
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_l = p.add_run(f'{label}: ')
    run_l.bold = True
    run_l.font.name = 'Arial'
    run_l.font.size = Pt(11)
    run_l.font.color.rgb = NAVY
    run_v = p.add_run(val)
    run_v.font.name = 'Arial'
    run_v.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(5)

add_hr(doc)

# ── STEP 3: THE SCRIPT ────────────────────────────────────────────────────────
add_section_heading(doc, 'STEP 3 — THE SCRIPT', NAVY)

# HOOK
add_label(doc, '[HOOK — 0 to 5 seconds]')
add_arabic(doc,
    'انت قفشت حاجة في بريطانيا؟ '
    'لو وقفت قدام حد في الـ queue… '
    'هو مش هيقولك حاجة. '
    'بس بعينيه هيقولك: أنت مش من هنا.',
    timing='0:00–0:05'
)
add_english(doc, '[On screen text overlay: THE QUEUE 🇬🇧 — explained]', colour=GOLD)

# CONTEXT SETUP
add_label(doc, '[CONTEXT SETUP — 5 to 20 seconds]')
add_arabic(doc,
    'الـ "queue" في بريطانيا مش بس طابور. '
    'ده دين. ده أخلاق. ده شخصية.\n\n'
    'الإنجليزي اللي بيستنى دوره في البقالة أو عند الأتوبيس '
    'مش بيعمل كده عشان عنده صبر. '
    'هو بيعمل كده عشان ده اللي بيثبت إنه إنسان محترم.\n\n'
    'وعشان تفهم ليه، محتاج ترجع لحاجة حصلت من أكتر من ٨٠ سنة.',
    timing='0:05–0:20'
)

# THE REAL STORY
add_label(doc, '[THE REAL STORY — 20 to 55 seconds]')
add_arabic(doc,
    'سنة ١٩٤٠. لندن تحت القصف. الأكل بقى بالبطاقة التموينية.\n\n'
    'كل أسبوع، المرأة الإنجليزية بتاخد بطاقتها وتقف في طابور — '
    'أحياناً ساعتين في البرد — قدام الجزار أو عند بايع السمك. '
    'مفيش تفضيل. مفيش واسطة. مفيش حد يتقدم.\n\n'
    'الحكومة وقتها عملت بوسترات تقول:\n'
    '"Queue and be fair" — وقف في دورك، ده معناه إنك إنسان شريف.\n\n'
    'يعني الـ queue مكانش بس طريقة تنظيم — ده بقى رمز وطني.\n'
    'اللي بيـ "queue jump" — اللي بيقطع الطابور — '
    'ده زي ما قال للناس: أنا أهم منكم.\n\n'
    'وده لحد دلوقتي أكبر إهانة اجتماعية في بريطانيا.\n\n'
    '— — —\n\n'
    'دلوقتي فكّر معايا. \n'
    'عندنا في مصر عندنا حاجة تانية خالص. \n'
    'عندنا الواسطة.\n\n'
    'مش فساد. مش غلط. '
    'ده إيه؟ ده إنك تعرف حد، يعرف حد، وتخلص شغلك.\n'
    'ده ذكاء اجتماعي. ده مجتمع بيشتغل بالعلاقات.\n\n'
    'الإنجليزي اخترع نظام: كلنا نستنى بالترتيب، والنظام ينظّم.\n'
    'المصري اخترع نظام تاني: أنا أعرف حد يساعدني، والعلاقة تنظّم.\n\n'
    'مفيش غلط ومفيش صح. في فلسفتين في الحياة.',
    timing='0:20–0:55'
)

# TAKEAWAY
add_label(doc, '[THE TAKEAWAY — 55 to 75 seconds]')
add_arabic(doc,
    'الدرس:\n'
    'لما الإنجليزي بيبص لك بعين زعلانة في الطابور — '
    'مش بيشوفك "أجنبي عديم الأدب." '
    'هو شايف حد بيكسر قاعدة أخلاقية عمرها ٨٠ سنة.\n\n'
    '📌 لو حصلك كده:\n'
    'قول بهدوء: "Sorry, I didn\'t realise — after you." \n'
    'جملة واحدة تمسح الموقف وتخليك إنت المحترم.\n\n'
    '📌 في الشغل:\n'
    'لو بتانتظر دورك في اجتماع أو email chain، '
    'استنى دورك. الإنجليزي بيلاحظ مين بيحترم الـ "turn-taking" ومين لأ.\n'
    'ده مش ضعف. ده ذكاء.',
    timing='0:55–1:15'
)

# CALL TO ACTION
add_label(doc, '[CALL TO ACTION — 1:15 to 1:30]')
add_arabic(doc,
    'إيه أغرب موقف حصل معاك في طابور هنا في بريطانيا؟\n'
    'ولا لما حد قطع طابورك وما اتكلمتيش؟ 😅\n\n'
    'اكتبهولي في الكومنتس — هعمله رييل.',
    timing='1:15–1:30'
)

add_hr(doc)

# ── CAPTION ───────────────────────────────────────────────────────────────────
add_section_heading(doc, 'CAPTION — Instagram & TikTok', NAVY)

add_label(doc, 'Arabic hook lines (lines 1–2):')
add_arabic(doc,
    'لو وقفت قدام حد في الطابور في بريطانيا، '
    'أنت مش بس كسرت قاعدة أدب — أنت كسرت قانون عمره ٨٠ سنة. 🇬🇧\n'
    'وعشان تفهم ليه، محتاج تعرف حكاية الحرب العالمية التانية وبطاقة التموين.'
)

add_label(doc, 'English discoverability line (line 3):')
add_english(doc,
    'The British queue isn\'t about manners — it\'s a WWII moral code. '
    'Here\'s what every Arab professional in the UK needs to know. 🔍'
)

add_label(doc, 'Hashtags:')
add_english(doc,
    '#مجتمع_عربي_في_بريطانيا #مصريين_في_بريطانيا #أطلس_للإنسانية\n'
    '#BritishCulture #Queueing #UKLife #CulturalIntelligence\n'
    '#SWANA #ArabsInUK #AtlasForHumanity\n'
    '#GrowingUpBritish #BritishHumour #DiasporaLife'
)

add_hr(doc)

# ── STEP 4: PRODUCTION NOTES ──────────────────────────────────────────────────
add_section_heading(doc, 'STEP 4 — PRODUCTION NOTES', NAVY)

add_english(doc, 'Visual direction:', bold=True)
add_english(doc, (
    '• Setting: Casual but put-together. Stand or sit slightly off-centre against a clean wall '
    '(cream or charcoal background preferred for brand colours).\n'
    '• Outfit: Smart-casual — blazer or structured top. Something you\'d wear to a London coffee '
    'shop. Warm but professional.\n'
    '• Camera: Eye-level, tight mid-shot (chest to top of head). No busy backgrounds.\n'
    '• Energy: Warm storytelling mode, not lecture mode. Use your hands — this is a conversation.'
))

add_english(doc, 'Text overlays (on screen):', bold=True)
add_english(doc, (
    '• 0:00 — "THE QUEUE 🇬🇧" (large, gold, centred)\n'
    '• 0:08 — "queue jumping = social crime" (small, bottom third, white)\n'
    '• 0:22 — "London, 1940 🗓" (small label, top left)\n'
    '• 0:38 — ""Queue and be fair" — UK WWII poster" (quoted text, small)\n'
    '• 0:48 — "الواسطة vs The Queue 🤔" (centred, Arabic + English, mid-screen)\n'
    '• 1:00 — \'Say: "Sorry — after you." ✅\' (tip card format, bottom third)\n'
    '• 1:10 — "turn-taking = professional capital 💼" (second tip card)'
))

add_english(doc, 'B-roll suggestions (optional):', bold=True)
add_english(doc, (
    '• Stock or creative-commons footage of a British queue (supermarket, bus stop, post office)\n'
    '• A quick phone-filmed clip of an actual UK queue if you can catch one this week\n'
    '• Optional: show a WWII ration book (prop or image) on screen briefly at 0:22'
))

add_english(doc, 'Photograph resource:', bold=True)
add_english(doc, (
    'URL: https://commons.wikimedia.org/wiki/Category:Rationing_in_the_United_Kingdom_in_World_War_II\n\n'
    'What it shows: Women and children queuing outside a London fishmonger during wartime rationing, '
    'circa 1940–1945. Black and white photograph, public domain via Wikimedia Commons.\n\n'
    'Caption credit line (if used in reel or caption):\n'
    'Photo: Wikimedia Commons / Imperial War Museum — Public Domain'
))

add_hr(doc)

# ── RESOURCES ─────────────────────────────────────────────────────────────────
add_section_heading(doc, 'RESOURCES', NAVY)

resources = [
    ('Wikimedia Commons — WWII UK Rationing photographs',
     'https://commons.wikimedia.org/wiki/Category:Rationing_in_the_United_Kingdom_in_World_War_II'),
    ('Debrett\'s — A Very British Queue',
     'https://debretts.com/a-very-british-queue/'),
    ('BBC Magazine — Is queuing really British?',
     'https://feeds.bbci.co.uk/news/magazine-23087024'),
    ('Parsons Green Prep — Queuing: A Great British Tradition',
     'https://www.parsonsgreenprep.co.uk/2025/03/queuing-a-great-british-tradition-that-fosters-a-sense-of-community/'),
    ('Holloway Express — Get in line: British culture of queuing',
     'https://hollowayexpress.org.uk/get-in-line-the-british-culture-of-queuing/'),
    ('The National — Ask Ali: Arabs and queue jumping',
     'https://www.thenationalnews.com/arts/ask-ali-arabs-parental-affection-and-queue-jumping-1.370819'),
    ('Scoop Empire — 12 things about Egyptians that give foreigners culture shock',
     'https://scoopempire.com/12-things-egyptians-give-foreigners-culture-shock/'),
    ('Social Bee — Instagram trends June 2026',
     'https://socialbee.com/blog/instagram-trends/'),
    ('Startups.co.uk — Social media trends UK June 2026',
     'https://startups.co.uk/news/social-media-trends-june-2026/'),
]

for title, url in resources:
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_bullet = p.add_run('• ')
    run_bullet.font.name = 'Arial'
    run_bullet.font.size = Pt(11)
    run_bullet.font.color.rgb = GOLD
    run_title = p.add_run(f'{title} — ')
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(11)
    run_url = p.add_run(url)
    run_url.font.name = 'Arial'
    run_url.font.size = Pt(10)
    run_url.font.color.rgb = RGBColor(0x1A, 0x5C, 0xA8)
    p.paragraph_format.space_after = Pt(5)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = '/home/user/Atlas-For-Humanity/Atlas_Reel_2026-06-16_TheQueue.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
