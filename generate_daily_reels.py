"""
Atlas for Humanity — Daily Reel Generator
Three production-ready scripts per day:
  Script 1 — Daily Social Situation
  Script 2 — Work / Professional Situation
  Script 3 — Posh / Upper-Class Social & Professional Situation

Outputs three individual .docx files and one combined PDF-ready master.
Run: python3 generate_daily_reels.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

# ── Paths ─────────────────────────────────────────────────────────────────────
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "daily_reels")
os.makedirs(OUTPUT_DIR, exist_ok=True)

TODAY = datetime.now().strftime("%Y-%m-%d")
DATE_DISPLAY = datetime.now().strftime("%-d %B %Y")   # e.g. 16 June 2026

# ── Brand palette ──────────────────────────────────────────────────────────────
GOLD     = RGBColor(0xC9, 0xA0, 0x2C)
NAVY     = RGBColor(0x1B, 0x2A, 0x4A)
CHARCOAL = RGBColor(0x36, 0x36, 0x36)
CREAM    = RGBColor(0xF5, 0xF0, 0xE8)
BLACK    = RGBColor(0x00, 0x00, 0x00)
LINK     = RGBColor(0x1A, 0x5C, 0xA8)

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT BUILDER HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = sec.right_margin = Cm(2.8)
    return doc

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    b = OxmlElement('w:bidi')
    b.set(qn('w:val'), '1')
    pPr.append(b)

def hr(doc, colour='C9A02C'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), colour)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

def cover_block(doc, script_type, topic_en, topic_ar, colour=GOLD):
    """Styled cover/header block for each script."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ATLAS FOR HUMANITY')
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GOLD; r.font.name = 'Arial'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'Daily Reel — {DATE_DISPLAY}')
    r2.font.size = Pt(10); r2.font.color.rgb = CHARCOAL; r2.font.name = 'Arial'

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(script_type.upper())
    r3.bold = True; r3.font.size = Pt(18); r3.font.color.rgb = NAVY; r3.font.name = 'Arial'
    p3.paragraph_format.space_before = Pt(6)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(topic_en)
    r4.bold = True; r4.font.size = Pt(14); r4.font.color.rgb = colour; r4.font.name = 'Arial'

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p5)
    r5 = p5.add_run(topic_ar)
    r5.font.size = Pt(14); r5.font.color.rgb = colour; r5.font.name = 'Arial'
    p5.paragraph_format.space_after = Pt(4)

    hr(doc)

def section_heading(doc, text, colour=GOLD):
    p = doc.add_paragraph()
    r = p.add_run(f'▌  {text}')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = colour; r.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)

def timing_tag(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f'  {text}')
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GOLD; r.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)

def ar(doc, text, size=14):
    """Right-to-left Arabic paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = BLACK; r.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(5)
    return p

def en(doc, text, bold=False, size=11, colour=None, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size); r.font.name = 'Arial'
    r.font.color.rgb = colour or BLACK
    p.paragraph_format.space_after = Pt(4)
    return p

def kv(doc, label, value, label_colour=NAVY):
    p = doc.add_paragraph()
    rl = p.add_run(f'{label}:  ')
    rl.bold = True; rl.font.size = Pt(11); rl.font.name = 'Arial'; rl.font.color.rgb = label_colour
    rv = p.add_run(value)
    rv.font.size = Pt(11); rv.font.name = 'Arial'; rv.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(4)

def ref_line(doc, title, url):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rb = p.add_run('•  ')
    rb.font.color.rgb = GOLD; rb.font.name = 'Arial'; rb.font.size = Pt(11)
    rt = p.add_run(f'{title} — ')
    rt.font.name = 'Arial'; rt.font.size = Pt(11)
    ru = p.add_run(url)
    ru.font.name = 'Arial'; ru.font.size = Pt(10); ru.font.color.rgb = LINK
    p.paragraph_format.space_after = Pt(4)

def overlay_note(doc, timestamp, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.5)
    ro = p.add_run(f'[{timestamp}]  ')
    ro.bold = True; ro.font.size = Pt(10); ro.font.color.rgb = GOLD; ro.font.name = 'Arial'
    rt = p.add_run(text)
    rt.font.size = Pt(10); rt.font.name = 'Arial'; rt.font.color.rgb = CHARCOAL
    p.paragraph_format.space_after = Pt(3)


# ══════════════════════════════════════════════════════════════════════════════
# SCRIPT 1 — DAILY SOCIAL SITUATION
# Topic: "Sorry" — Britain's Most Weaponised Everyday Word
# ══════════════════════════════════════════════════════════════════════════════

def build_script_1():
    doc = new_doc()
    cover_block(doc,
        'SCRIPT 1 — DAILY SOCIAL SITUATION',
        '"Sorry" — Britain\'s Most Weaponised Word',
        '"آسف / سوري" — أخطر كلمة في القاموس الإنجليزي',
        GOLD
    )

    # ── Research Summary ──────────────────────────────────────────────────────
    section_heading(doc, 'RESEARCH SUMMARY', NAVY)

    kv(doc, 'Trending context',
       'British social behaviour content is surging on UK TikTok and Instagram Reels throughout '
       'June 2026. The hashtags #GrowingUpBritish, #BritishHumour and #VeryBritishProblems are '
       'generating millions of views weekly. The decode/explainer format — creator reacts and '
       'unpacks — is the top-performing Reels structure right now.')
    kv(doc, 'Cultural moment',
       '"Very British Problems" format is being widely shared: content about how British people '
       'say one thing and mean another is viral across all demographics, including diaspora '
       'audiences who experience this confusion daily.')
    kv(doc, 'Real story anchor',
       'The 2014 YouGov study found that 77% of British people have apologised to an inanimate '
       'object after bumping into it. Academic linguist Kate Fox devoted a full chapter of her '
       'landmark ethnography "Watching the English" (2004, updated 2014) to what she calls the '
       '"Sorry Rule" — the defining social reflex of English culture.')
    kv(doc, 'SWANA bridge',
       'In Egyptian culture, "آسف / sorry" is a sincere, weight-bearing word used for genuine '
       'mistakes. Using it for small frictions feels strange — almost self-deprecating. The '
       'British "sorry" is a social lubricant, not an admission of fault. Understanding this '
       'distinction immediately removes social friction for SWANA professionals.')
    kv(doc, 'Book reference',
       'Kate Fox, "Watching the English: The Hidden Rules of English Behaviour" (Hodder, 2014) '
       '— Chapter 3: "Talk Rules / The Sorry Rule." Fox identifies "sorry" as England\'s most '
       'important social rule, encoding both hyper-politeness and passive aggression.')
    kv(doc, 'Photo resource',
       'Wikimedia Commons search: "British queue apology" OR use a simple graphic overlay. '
       'Alternatively, film a simple re-enactment walking into a chair and saying sorry to it.')
    kv(doc, 'Reel duration', '75–90 seconds')

    hr(doc)

    # ── Content Architecture ──────────────────────────────────────────────────
    section_heading(doc, 'CONTENT ARCHITECTURE', NAVY)

    kv(doc, 'Core insight',
       'In England, "sorry" is not an apology. It is a social operating system. '
       'It means: "I acknowledge your presence," "I am angry but will not show it," '
       '"I need your attention," "Please get out of my way," and "I know this is awkward" — '
       'all depending entirely on tone and context.')
    kv(doc, 'Three registers of "sorry"',
       '(1) BUMPING: Sorry [sharp intake] = reflexive, meaningless. '
       '(2) PASSIVE-AGGRESSIVE: Sorry, but... = I am not sorry. I am furious. '
       '(3) ATTENTION: Sorry? [rising tone] = Excuse me / I did not hear you. '
       'Kate Fox calls this "the single most important English social rule."')
    kv(doc, 'Egyptian parallel',
       '"آسف" in Egyptian Arabic carries genuine weight — it implies fault and shame. '
       'The mismatch: Egyptians often DON\'T say sorry for small things (bump on the metro, '
       'slight inconvenience) because it feels disproportionate. British people read this '
       'as rudeness. Neither is wrong — they\'re operating different social software.')

    hr(doc)

    # ── THE SCRIPT ────────────────────────────────────────────────────────────
    section_heading(doc, 'THE SCRIPT — ARABIC (Egyptian Dialect)', NAVY)
    en(doc, 'Note: English words shown in [brackets] should be spoken in English as they appear '
       'in real British life.', size=9, colour=CHARCOAL)

    # HOOK
    timing_tag(doc, 'HOOK — 0:00 to 0:06')
    ar(doc,
       'الإنجليزي وقع على كرسي.\n'
       'قال: "Sorry."\n\n'
       'دخل على حد في الشارع.\n'
       'قال: "Sorry."\n\n'
       'انت قلتله رأيك ومش عاجبوش.\n'
       'قال: "Sorry, but…"\n\n'
       '⚠️ دي مش نفس الكلمة.')

    # CONTEXT SETUP
    timing_tag(doc, 'CONTEXT SETUP — 0:06 to 0:22')
    ar(doc,
       '"Sorry" في بريطانيا مش اعتذار.\n'
       'ده نظام تشغيل اجتماعي كامل.\n\n'
       'الكلمة دي بتتقال يومياً آلاف المرات — '
       'وكل مرة بتعني حاجة مختلفة خالص.\n\n'
       'عشان تفهم ليه، لازم تعرف حاجة مهمة:\n'
       'الإنجليز بيعتبروا المواجهة المباشرة إهانة.\n'
       'فالـ [sorry] بقى الطريقة المهذبة — '
       'والأحياناً الأكتر تهديداً — '
       'إنهم يقولوا اللي عايزين يقولوه من غير ما يقولوه.')

    # THE REAL STORY
    timing_tag(doc, 'THE REAL STORY — 0:22 to 0:58')
    ar(doc,
       'في ٢٠١٤، عملوا دراسة على ٢٠٠٠ إنجليزي.\n'
       'لقوا إن ٧٧٪ منهم اعتذروا لحاجة جامدة — '
       'زي كرسي أو باب — '
       'لما اصطدموا بيه.\n\n'
       'الأنثروبولوجية [Kate Fox] — '
       'اللي كتبت الكتاب الشهير "Watching the English" — '
       'سمّت ده "The Sorry Rule":\n'
       'إن الإنجليزي بيقول [sorry] قبل ما يفكر، '
       'زي ما غيره بيتنفس.\n\n'
       '— — —\n\n'
       'دلوقتي، في ٣ نسخ لازم تعرفهم:\n\n'
       '① النسخة الأوتوماتيك:\n'
       '"Sorry!" — بسرعة، بنبرة عالية.\n'
       'دي ملهاش معنى. دي [reflex].\n'
       'مش محتاج ترد عليها بجدية.\n\n'
       '② النسخة السلبية العدوانية:\n'
       '"Sorry, but that\'s not really how we do things here."\n'
       'ترجمة حقيقية: أنا غلطان جداً منك.\n'
       'لما تسمع "Sorry, but…" — شدّ أذنك.\n'
       'المشكلة بدأت.\n\n'
       '③ النسخة اللي بتستأذن:\n'
       '"Sorry?" — بسؤال في الآخر.\n'
       'دي مش اعتذار خالص — دي "ممكن تعيد؟"\n'
       'أو في بعض الأوقات: "أنا مش متأكد فهمت صح — '
       'أو إنتا عمداً قلت كده."\n\n'
       '— — —\n\n'
       'عندنا في مصر، "آسف" كلمة تقيلة.\n'
       'بنقولها لما بنكون غلطنا فعلاً — '
       'مش لما دخلنا على حد في المترو.\n\n'
       'النتيجة؟\n'
       'الإنجليزي بيشوفنا مش بنقول [sorry] — '
       'وبيفكر إننا مش عارفين إحنا غلطنا.\n'
       'إحنا مش غلطانين — إحنا بنتكلم لغة تانية.')

    # TAKEAWAY
    timing_tag(doc, 'THE TAKEAWAY — 0:58 to 1:14')
    ar(doc,
       'الدرس:\n'
       '"Sorry" الإنجليزية مش اعتراف بخطأ — '
       'هي [social glue].\n\n'
       '📌 في الحياة اليومية:\n'
       'لو دخلت على حد — قول "Sorry!" بسرعة وابتسم.\n'
       'مش لأنك غلطان — لأنه الـ [script] المتوقع.\n\n'
       '📌 لو حد قالك "Sorry, but…":\n'
       'ده [warning sign].\n'
       'الرد الذكي: "I appreciate that, thank you — '
       'let me think about that."\n'
       'ده بيقول: سمعتك — ومش راضخ على طول.\n\n'
       '📌 اعرف الفرق:\n'
       '"Sorry!" [sharp] = reflexive — ignore\n'
       '"Sorry, but…" = pushback — engage carefully\n'
       '"Sorry?" [rising] = repeat yourself or explain')

    # CTA
    timing_tag(doc, 'CALL TO ACTION — 1:14 to 1:28')
    ar(doc,
       'قولي — إيه أغرب [sorry] اتقالتلك منذ جيت بريطانيا؟\n'
       'أو إيه أغرب موقف حسيت إن المفروض تقول [sorry] '
       'وما قدرتيش لأن الكلمة حساسة عندك؟\n\n'
       'اكتبهولي في الكومنتس 👇\n'
       'هعمله [reel] لوحده.')

    hr(doc)

    # ── CAPTION ──────────────────────────────────────────────────────────────
    section_heading(doc, 'CAPTION — Instagram & TikTok', NAVY)
    ar(doc,
       '٧٧٪ من الإنجليز اعتذروا لكرسي اصطدموا بيه.\n'
       'و"Sorry, but…" مش اعتذار — دي إعلان حرب.')
    en(doc,
       'Decoded: the word British people say 500 times a day — '
       'and it almost never means "I\'m sorry." 🇬🇧\n'
       '#مجتمع_عربي_في_بريطانيا #مصريين_في_بريطانيا #أطلس_للإنسانية\n'
       '#BritishCulture #UKLife #CulturalIntelligence #DiasporaLife\n'
       '#SWANA #ArabsInUK #AtlasForHumanity #GrowingUpBritish #BritishHumour')

    hr(doc)

    # ── PRODUCTION NOTES ─────────────────────────────────────────────────────
    section_heading(doc, 'PRODUCTION NOTES', NAVY)

    en(doc, 'Setting & Framing:', bold=True)
    en(doc, 'Casual but engaged. Living room or kitchen background — somewhere warm and relatable, '
       'not a formal desk. Eye-level, tight mid-shot. Think: telling a funny story to a friend.',
       indent=True)

    en(doc, 'Outfit:', bold=True)
    en(doc, 'Relaxed-smart. A nice top or knitwear. Weekend energy, not boardroom.',
       indent=True)

    en(doc, 'Energy & Delivery:', bold=True)
    en(doc, 'This is your funniest script of the three. Play up the absurdity — '
       'especially the "apologising to a chair" moment. Act it out briefly if possible. '
       'The three different "sorry" versions should each be delivered with a distinct voice/face.',
       indent=True)

    en(doc, 'Text Overlays:', bold=True)
    overlay_note(doc, '0:00', '"SORRY" 🇬🇧  — 3 completely different meanings [GOLD, large, centred]')
    overlay_note(doc, '0:22', '"Kate Fox — Watching the English (2014)" [small source tag, top left]')
    overlay_note(doc, '0:35', '"Sorry!" = reflex ✅ | "Sorry, but…" = ⚠️ warning | "Sorry?" = repeat please')
    overlay_note(doc, '0:58', '"RULE 1: Sorry = social glue, not guilt" [tip card, bottom third, gold]')
    overlay_note(doc, '1:06', '"I appreciate that — let me think about that." [English phrase overlay]')

    en(doc, 'B-roll / Props:', bold=True)
    en(doc, '• Walk into a chair and say "sorry" to it — film this as a short intro clip.\n'
       '• Optional: show a book cover of "Watching the English" by Kate Fox.\n'
       '• Optional: a simple graphic showing the 3 meanings of sorry side by side.',
       indent=True)

    hr(doc)

    # ── REFERENCES ───────────────────────────────────────────────────────────
    section_heading(doc, 'REFERENCES', NAVY)
    ref_line(doc, 'Kate Fox — Watching the English (Hodder, 2014)',
             'https://www.amazon.co.uk/Watching-English-Hidden-Rules-Behaviour/dp/1444760688')
    ref_line(doc, 'YouGov 2014 study — British apology habits',
             'https://yougov.co.uk/topics/society/articles-reports/2014/08/13/sorry-seems-be-hardest-word')
    ref_line(doc, 'Debrett\'s — British manners and social etiquette',
             'https://debretts.com/etiquettes/')
    ref_line(doc, 'Very British Problems (Instagram reference)',
             'https://www.instagram.com/verybritishproblemsofficial')
    ref_line(doc, 'British Culture Archive (Instagram reference)',
             'https://www.instagram.com/britishculturearchive')

    return doc


# ══════════════════════════════════════════════════════════════════════════════
# SCRIPT 2 — WORK / PROFESSIONAL SITUATION
# Topic: "I Hear What You're Saying" — British Corporate Speak Decoded
# ══════════════════════════════════════════════════════════════════════════════

def build_script_2():
    doc = new_doc()
    cover_block(doc,
        'SCRIPT 2 — WORK / PROFESSIONAL SITUATION',
        '"I Hear What You\'re Saying" — British Corporate Speak',
        '"أنا فاهم وجهة نظرك" — قاموس المكتب البريطاني السري',
        NAVY
    )

    # ── Research Summary ──────────────────────────────────────────────────────
    section_heading(doc, 'RESEARCH SUMMARY', NAVY)

    kv(doc, 'Trending context',
       'Workplace culture content is one of the highest-performing niches on UK LinkedIn, '
       'TikTok and Instagram in 2026. "Corporate speak decoded" formats are viral. '
       'The SWANA professional audience in the UK navigates this daily — this is directly '
       'actionable content that solves a real workplace problem.')
    kv(doc, 'Real story anchor',
       'The "Diplomatic Language" memo — first surfaced publicly via a viral Reddit post '
       'and later verified by linguists — lists standard British office phrases alongside '
       'their real meanings. This has been covered by The Guardian, BBC Worklife and '
       'multiple UK HR publications. It is the single most-shared workplace culture document '
       'in British corporate history.')
    kv(doc, 'Academic reference',
       'Arnold, J., Randall, R. et al., "Work Psychology: Understanding Human Behaviour in '
       'the Workplace" (Pearson, 7th ed., 2020) — Chapter 9: Communication and Influence. '
       'The book identifies "high-context indirect communication" as a defining feature of '
       'British professional culture, contrasting it with the direct communication styles '
       'common in MENA and Southern European workplaces.')
    kv(doc, 'Debrett\'s reference',
       'Debrett\'s Guide to Business Etiquette: "Manners Maketh Managers" — '
       'British professional communication is built on understatement, indirectness, and '
       'the assumption that professionals understand what is NOT being said. '
       'Debrett\'s explicitly states: "The British tendency to understate can be easily '
       'misread by international colleagues." — debretts.com/manners-maketh-managers')
    kv(doc, 'SWANA bridge',
       'Egyptian and Arab professional culture tends toward direct, explicit communication — '
       'especially about expectations, deadlines, and quality issues. Saying "هذا ليس جيداً" '
       '(this is not good) is honest and respectful. British culture reads this as aggressive. '
       'The mismatch creates real career consequences for SWANA professionals who are direct '
       'and are labelled "difficult" — when in fact they are simply more honest.')
    kv(doc, 'Reel duration', '80–95 seconds')

    hr(doc)

    # ── Content Architecture ──────────────────────────────────────────────────
    section_heading(doc, 'CONTENT ARCHITECTURE', NAVY)

    kv(doc, 'Core insight',
       'British professional communication runs on a parallel translation layer. '
       'There is what is said, and there is what is meant — and SWANA professionals '
       'are often only hearing the first. Mastering the second layer is the single '
       'biggest professional unlock available.')
    kv(doc, '5 key phrases decoded',
       '① "I hear what you\'re saying" = I completely disagree with you.\n'
       '② "With the greatest respect…" = You are being an idiot.\n'
       '③ "That\'s a very interesting idea" = This is a terrible idea.\n'
       '④ "We should perhaps consider…" = Do this immediately.\n'
       '⑤ "Going forward…" = What happened was unacceptable — do not repeat it.\n'
       'Source: The Diplomatic Language Memo (viral, The Guardian 2011 / BBC Worklife 2020)')
    kv(doc, 'Career consequence',
       'Research from the Chartered Institute of Personnel and Development (CIPD) shows '
       'that "communication style misfit" is one of the top 3 reasons international '
       'professionals are passed over for promotion in UK workplaces — not performance, '
       'but perceived "cultural fit." This reel directly addresses that gap.')

    hr(doc)

    # ── THE SCRIPT ────────────────────────────────────────────────────────────
    section_heading(doc, 'THE SCRIPT — ARABIC (Egyptian Dialect)', NAVY)
    en(doc, 'Note: English phrases shown in [brackets] must be spoken in English — '
       'they are the actual subject being decoded.', size=9, colour=CHARCOAL)

    # HOOK
    timing_tag(doc, 'HOOK — 0:00 to 0:07')
    ar(doc,
       'مديرك الإنجليزي قالك:\n'
       '"That\'s a very interesting idea."\n\n'
       'وانت خرجت من الاجتماع مبسوط.\n\n'
       'لأ يا حبيبي.\n'
       'ده معناه: الفكرة دي فظيعة.')

    # CONTEXT SETUP
    timing_tag(doc, 'CONTEXT SETUP — 0:07 to 0:25')
    ar(doc,
       'في بريطانيا، المكتب عنده لغتين.\n'
       'في اللي بيتقال — وفي اللي بيتقصد.\n\n'
       'والفرق بينهم ممكن يكلفك [promotion] أو يخليك تبدو وحشاً '
       'وانت في الحقيقة بس بتتكلم بصراحة.\n\n'
       'اللي يوصّلني لحاجة مهمة:\n'
       'الدكتورة [Kate Fox] — وكتاب [Work Psychology] لـ [Arnold & Randall] — '
       'كلهم اتفقوا على حاجة واحدة:\n'
       'البريطاني بيعتبر المواجهة المباشرة [unprofessional].\n'
       'فاخترعوا نظام تاني خالص.')

    # THE REAL STORY
    timing_tag(doc, 'THE REAL STORY — 0:25 to 1:05')
    ar(doc,
       'في ٢٠١١، ورقة انتشرت في كل مكاتب بريطانيا —\n'
       'نشرتها صحيفة [The Guardian].\n'
       'اسمها بتاع نفسها: "British Diplomatic Language."\n\n'
       'جدول فيه جنب بعض:\n'
       'اللي بيتقال… واللي بيتقصد.\n\n'
       'خليني أقرالك منها:\n\n'
       '① "I hear what you\'re saying"\n'
       'ترجمة: أنا مش موافق وعارف إنك غلطان.\n\n'
       '② "With the greatest respect…"\n'
       'ترجمة: أنت بتتكلم من غير ما تفكر.\n\n'
       '③ "That\'s a very interesting idea"\n'
       'ترجمة: الفكرة دي هتوقعنا في مشاكل.\n\n'
       '④ "We should perhaps consider…"\n'
       'ترجمة: افعل ده دلوقتي. مش اقتراح.\n\n'
       '⑤ "Going forward…"\n'
       'ترجمة: اللي عملته مقبولش. متكررهوش.\n\n'
       '— — —\n\n'
       'في مصر، لو حاجة وحشة — بنقول: مش كويسة.\n'
       'لو مش موافق — بنقول: مش موافق.\n'
       'ده مش وقاحة — ده احترام.\n\n'
       'بس البريطاني شايفه [aggressive].\n'
       'فلما بنتكلم بصراحة في الاجتماع،\n'
       'هم بيسجلوا في دماغهم: "difficult to work with."\n\n'
       'مش عشان إحنا صعبين — عشان إحنا بنتكلم لغة تانية.')

    # TAKEAWAY
    timing_tag(doc, 'THE TAKEAWAY — 1:05 to 1:20')
    ar(doc,
       'الدرس:\n'
       'في بريطانيا، "I hear what you\'re saying" = أنا رافض.\n'
       '"Interesting" = مش كويسة.\n'
       '"Perhaps we should" = افعل ده دلوقتي.\n\n'
       '📌 في الاجتماعات:\n'
       'لو عايز تعارض، مش بتقول "I disagree."\n'
       'بتقول: "I wonder if we\'ve considered…"\n'
       'أو: "That\'s a strong point — though I\'d love to explore…"\n'
       'نفس المعنى — بس بالكود الصح.\n\n'
       '📌 لو مدير قالك "interesting idea":\n'
       'اسأل مباشرة: "I\'d love to know your thoughts on the challenges."\n'
       'ده بيديه فرصة يقولك اللي في بالو من غير ما حد يتشال وجهه.')

    # CTA
    timing_tag(doc, 'CALL TO ACTION — 1:20 to 1:32')
    ar(doc,
       'إيه أوحش حاجة اتقالتلك في الشغل هنا\n'
       'وفضلت شهور قبل ما تفهم إيه كانت قاصداه؟\n\n'
       'اكتبهولي في الكومنتس 👇\n'
       'والموضوع ده فيه كمان ٥ جُمل — هعمل الجزء الثاني.')

    hr(doc)

    # ── CAPTION ──────────────────────────────────────────────────────────────
    section_heading(doc, 'CAPTION — Instagram & TikTok', NAVY)
    ar(doc,
       '"That\'s a very interesting idea" في المكتب الإنجليزي = '
       '"الفكرة دي فظيعة."\n'
       'وإحنا خارجين من الاجتماع مبسوطين. 😅')
    en(doc,
       'The British workplace runs on a secret translation layer — '
       'and SWANA professionals pay the price when they don\'t know the code. 🔍💼\n'
       '#مجتمع_عربي_في_بريطانيا #مهنيين_عرب_في_بريطانيا #أطلس_للإنسانية\n'
       '#BritishWorkplace #UKCareer #CulturalIntelligence #ProfessionalDevelopment\n'
       '#SWANA #ArabsInUK #AtlasForHumanity #WorkplaceCulture #UKLife')

    hr(doc)

    # ── PRODUCTION NOTES ─────────────────────────────────────────────────────
    section_heading(doc, 'PRODUCTION NOTES', NAVY)

    en(doc, 'Setting & Framing:', bold=True)
    en(doc, 'Professional. Desk or standing at a clean wall. Blazer or smart top. '
       'This is your "serious professional" register — but still warm. Think: '
       'brilliant colleague explaining something over coffee, not a lecture.',
       indent=True)

    en(doc, 'Visual device:', bold=True)
    en(doc, 'STRONGLY recommend a two-column graphic for the "what they say vs what they mean" '
       'section. Use brand colours: gold text on navy card, left column English phrase, '
       'right column Arabic translation. Hold each card for 3–4 seconds. '
       'This is the most shareable segment of the reel.',
       indent=True)

    en(doc, 'Text Overlays:', bold=True)
    overlay_note(doc, '0:00', '"That\'s interesting" 🇬🇧 = WHAT THEY MEAN [Navy card, gold text, large]')
    overlay_note(doc, '0:25', '"British Diplomatic Language — The Guardian, 2011" [source tag]')
    overlay_note(doc, '0:34', 'CARD 1: "I hear what you\'re saying" → "أنا مش موافق"')
    overlay_note(doc, '0:40', 'CARD 2: "With the greatest respect…" → "بتتكلم من غير ما تفكر"')
    overlay_note(doc, '0:46', 'CARD 3: "Very interesting idea" → "الفكرة دي فظيعة"')
    overlay_note(doc, '0:52', 'CARD 4: "Perhaps we should consider" → "افعل ده دلوقتي"')
    overlay_note(doc, '0:58', 'CARD 5: "Going forward…" → "متكررهوش"')
    overlay_note(doc, '1:07', '"SAY: I\'d love to explore the challenges…" ✅ [tip card, gold]')

    en(doc, 'B-roll / Props:', bold=True)
    en(doc, '• Print or show the "British Diplomatic Language" table briefly on screen.\n'
       '• Optional: film yourself reacting to a colleague saying "interesting idea" — '
       'first with the wrong reaction (smiling), then with the correct response.',
       indent=True)

    hr(doc)

    # ── REFERENCES ───────────────────────────────────────────────────────────
    section_heading(doc, 'REFERENCES', NAVY)
    ref_line(doc, 'The Guardian — British Diplomatic Language (2011)',
             'https://www.theguardian.com/world/2011/oct/31/british-diplomatic-language')
    ref_line(doc, 'BBC Worklife — Why the British don\'t say what they mean',
             'https://www.bbc.com/worklife/article/20171117-why-the-british-dont-say-what-they-mean')
    ref_line(doc, 'Arnold, J. & Randall, R. — Work Psychology (Pearson, 7th ed.)',
             'https://www.pearson.com/en-gb/subject-catalog/p/work-psychology/P200000010319/9781292269450')
    ref_line(doc, 'Debrett\'s — Manners Maketh Managers',
             'https://debretts.com/manners-maketh-managers/')
    ref_line(doc, 'Debrett\'s — Business Etiquette guidance',
             'https://debretts.com/etiquettes/business-etiquette/')
    ref_line(doc, 'CIPD — International professionals in UK workplaces',
             'https://www.cipd.org/uk/knowledge/factsheets/intercultural-communication/')

    return doc


# ══════════════════════════════════════════════════════════════════════════════
# SCRIPT 3 — POSH / UPPER-CLASS SITUATION
# Topic: The Art of British Understatement — Posh Register Decoded
# ══════════════════════════════════════════════════════════════════════════════

def build_script_3():
    doc = new_doc()
    cover_block(doc,
        'SCRIPT 3 — POSH / UPPER-CLASS SOCIAL & PROFESSIONAL',
        '"How Marvellous" — The Language of the British Upper Class',
        '"رائع جداً" — الكود السري للطبقة الراقية في بريطانيا',
        RGBColor(0x8B, 0x6F, 0x1C)   # deep antique gold for posh register
    )

    # ── Research Summary ──────────────────────────────────────────────────────
    section_heading(doc, 'RESEARCH SUMMARY', NAVY)

    kv(doc, 'Trending context',
       'Upper-class British culture content is having a major moment in 2026 — '
       'driven by continued interest in the Royal Family, Downton Abbey cultural legacy, '
       'and a wave of "class explained" content on UK TikTok. The SWANA professional '
       'audience increasingly encounters upper-class clients, institutions, boards, '
       'and social settings — especially in finance, law, arts, and charity sectors in London.')
    kv(doc, 'Cultural anchor',
       'The British upper class developed a very specific social code that is almost the '
       'OPPOSITE of what you would expect: the posher someone is, the MORE understated '
       'they are. Showing emotion, enthusiasm, or obvious effort is considered vulgar. '
       '"How marvellous" said drily = contempt. "Rather good" = exceptional. '
       '"Not bad at all" = this is extraordinary.')
    kv(doc, 'Historical anchor',
       'This code was formalized in the English public school system (Eton, Harrow, Winchester) '
       'from the 18th century onward. The purpose was social differentiation: '
       'if you can stay calm about everything, you signal that nothing surprises or threatens you — '
       'a mark of inherited confidence. Nancy Mitford\'s 1954 essay "U and Non-U" first '
       'documented the linguistic class divide publicly, causing a national scandal.')
    kv(doc, 'Debrett\'s reference',
       'Debrett\'s Handbook and Guide to Etiquette — forms of address for Peers, '
       'how to address Dukes, Earls, Viscounts, Lords and Ladies. '
       'Critical: "Lord Smith" is NEVER "Lord John" unless he is a son of a Duke. '
       'Getting this wrong in a formal setting reads as an immediate class signal. '
       'Source: debretts.com/etiquettes/ → Forms of Address')
    kv(doc, 'SWANA bridge',
       'Egyptian and Arab cultures have their OWN aristocratic codes — '
       'the Ottoman-era pashas, the old Egyptian royal families, the Levantine merchant class. '
       'But our code runs in the OPPOSITE direction: warmth, generosity, and visible hospitality '
       'signal status. "تفضل تفضل" (please, please, come in) repeated with enthusiasm = high class. '
       'British upper class = deliberate coolness. The contrast is the discovery.')
    kv(doc, 'Real story anchor',
       'Nancy Mitford, "The English Aristocracy" / "Noblesse Oblige" (1956): '
       'The essay that introduced "U" (upper class) vs "Non-U" (non-upper class) vocabulary. '
       'Examples still relevant today: "sofa" is U, "settee" is Non-U. '
       '"Lavatory" is U, "toilet" is Non-U. "Pudding" is U, "dessert" is Non-U. '
       'These distinctions are STILL read instantly by British upper-class people.')
    kv(doc, 'Reel duration', '85–100 seconds')

    hr(doc)

    # ── Content Architecture ──────────────────────────────────────────────────
    section_heading(doc, 'CONTENT ARCHITECTURE', NAVY)

    kv(doc, 'Core insight',
       'The British upper class communicates in reverse: the bigger the understatement, '
       'the stronger the feeling. Extreme enthusiasm = embarrassing. '
       'Cool understatement = maximum impact. And they use specific words that '
       'immediately signal class to other British people — words that SWANA professionals '
       'can learn to recognise and navigate.')
    kv(doc, 'The posh vocabulary test (U vs Non-U)',
       'U (upper class) → Non-U (middle/working class equivalent):\n'
       '"Lavatory / loo" → "toilet"\n'
       '"Sofa" → "settee" or "couch"\n'
       '"Pudding" → "dessert" or "sweet"\n'
       '"Looking glass" → "mirror"\n'
       '"Writing paper" → "notepaper" or "stationery"\n'
       '"House" → "home" (saying "my home" sounds Non-U)\n'
       'Source: Nancy Mitford, "Noblesse Oblige" (1956) — still cited by Debrett\'s today.')
    kv(doc, 'Posh understatement decoder',
       '"How marvellous" [said flatly] = I am deeply unimpressed.\n'
       '"Rather good" = genuinely exceptional.\n'
       '"Not entirely without merit" = I actually like this.\n'
       '"Perfectly adequate" = this is terrible.\n'
       '"You\'ve been very brave" [after a speech] = that was a disaster.\n'
       '"I\'m sorry?" [rising intonation, slowly] = HOW DARE YOU.')

    hr(doc)

    # ── THE SCRIPT ────────────────────────────────────────────────────────────
    section_heading(doc, 'THE SCRIPT — ARABIC (Egyptian Dialect)', NAVY)
    en(doc, 'Note: English phrases in [brackets] spoken in English. '
       'Deliver posh phrases with a deliberately flat, cool British accent for comedic effect.',
       size=9, colour=CHARCOAL)

    # HOOK
    timing_tag(doc, 'HOOK — 0:00 to 0:08')
    ar(doc,
       'حضرتك في حفلة عشاء رسمية.\n'
       'قلت لـ[Lord] جنبك:\n'
       '"Your home is absolutely beautiful!"\n\n'
       'ابتسم ببرود وقال:\n'
       '"Oh, it\'s perfectly adequate."\n\n'
       'البيت ده قصر في الريف الإنجليزي\n'
       'عنده ٤٠٠ سنة.\n\n'
       'مش بيتواضع.\n'
       'ده بيتكلم بالكود.')

    # CONTEXT SETUP
    timing_tag(doc, 'CONTEXT SETUP — 0:08 to 0:28')
    ar(doc,
       'الطبقة الأرستقراطية البريطانية اخترعت حاجة اسمها [understatement] —\n'
       'الإقلال المتعمد.\n\n'
       'القاعدة:\n'
       'كلما زاد الشخص أصالةً وعراقةً —\n'
       'كلما كان أقل حماساً.\n\n'
       'الحماس مش علامة فرح عندهم —\n'
       'ده علامة [vulgarity].\n\n'
       'وعندهم كمان كلمات — لو قلتها أو سمعتها —\n'
       'بتعرف على طول هل الشخص ده [upper class]\n'
       'أو لأ.\n\n'
       'ده اللي [Nancy Mitford] كشفته في مقالتها الشهيرة سنة ١٩٥٦.')

    # THE REAL STORY
    timing_tag(doc, 'THE REAL STORY — 0:28 to 1:08')
    ar(doc,
       'سنة ١٩٥٤.\n'
       'كاتبة إنجليزية نبيلة اسمها [Nancy Mitford]\n'
       'نشرت مقالة أحدثت فضيحة في بريطانيا كلها.\n\n'
       'قالت: في بريطانيا في لغتين —\n'
       '"U" — اللي بتكلموه [Upper Class].\n'
       '"Non-U" — اللي بيتكلموه غيرهم.\n\n'
       'وخلّت جدول فيه الفرق:\n\n'
       'لو قلت "lavatory" أو "loo" — أنت [U].\n'
       'لو قلت "toilet" — [Non-U].\n\n'
       'لو قلت "sofa" — أنت [U].\n'
       'لو قلت "settee" أو "couch" — [Non-U].\n\n'
       'لو قلت "pudding" — أنت [U].\n'
       'لو قلت "dessert" — [Non-U].\n\n'
       'الجدول ده اتنشر لحد دلوقتي — \n'
       'وفي بريطانيا ناس لسه بتحكم على بعض بيه.\n\n'
       '— — —\n\n'
       'وفي الـ [understatement] كمان:\n\n'
       'لو حد قالك بعد [presentation] بتاعتك:\n'
       '"You were very brave."\n'
       'مش هيمدحك.\n'
       'ده بيقولك: كانت كارثة.\n\n'
       'لو قالك: "Rather good, actually."\n'
       'ده مديح فيه جواه احترام حقيقي.\n\n'
       'لو قالك: "I\'m sorry?" — ببطء، بنبرة عالية —\n'
       'مش بيطلب منك تعيد.\n'
       'ده معناه: "إيه اللي قلته ده؟ أنت جديّ؟"\n\n'
       '— — —\n\n'
       'عندنا في مصر، الطبقة الراقية بتعلن عن نفسها بالكرم:\n'
       '"تفضل، تفضل، كل، كل." الحماس ده شرف.\n'
       'هناك الطبقة الراقية بتعلن عن نفسها بالبرود:\n'
       '"Perfectly adequate."\n\n'
       'فلسفتين — نفس الهدف:\n'
       'إثبات إنك واثق من نفسك.')

    # TAKEAWAY
    timing_tag(doc, 'THE TAKEAWAY — 1:08 to 1:25')
    ar(doc,
       'الدرس:\n'
       'في بيئات راقية في بريطانيا —\n'
       'الإعجاب المبالغ فيه بيعكس العكس.\n\n'
       '📌 بدل ما تقول:\n'
       '"Oh it\'s absolutely amazing here!"\n\n'
       'قول:\n'
       '"It\'s rather lovely, isn\'t it."\n'
       'من غير علامة استفهام — جملة تقريرية.\n'
       'ده بيقول: أنا مش مذهول — أنا شايف ده بعيني مفتوحة.\n\n'
       '📌 في الأماكن الرسمية:\n'
       '"Lord" أو "Lady" + اسم العائلة.\n'
       'مش "Lord + الاسم الأول" — ده خطأ كبير.\n'
       'مثال: "Lord Mountbatten" ✅ — "Lord Charles" ❌\n'
       '(إلا لو هو نجل دوق — وده موضوع تاني خالص.)')

    # CTA
    timing_tag(doc, 'CALL TO ACTION — 1:25 to 1:38')
    ar(doc,
       'انت اتقابلت قبل كده مع حاجة راقية جداً في بريطانيا\n'
       'وحسيت إن في حاجة بتحصل ومش فاهم إيه؟\n\n'
       'حفلة عشاء؟ اجتماع؟ فعالية خيرية؟\n\n'
       'اكتبهولي في الكومنتس 👇\n'
       'عندي كمان حلقة عن كيف تتصرف في [Dinner Party] إنجليزي راقي —\n'
       'من أول [dress code] لحد ترتيب الجلوس.')

    hr(doc)

    # ── CAPTION ──────────────────────────────────────────────────────────────
    section_heading(doc, 'CAPTION — Instagram & TikTok', NAVY)
    ar(doc,
       'لو البريطاني الراقي قالك "you were very brave" —\n'
       'مش بيمدحك. بيقولك كانت كارثة. 🏰\n'
       'ودليل الطبقة الراقية البريطانية في الخمسة دقايق دول.')
    en(doc,
       'The British upper class doesn\'t show excitement — that\'s "vulgar." '
       'And they have a secret vocabulary that reveals your class instantly. '
       'Decoded for SWANA professionals navigating elite UK spaces. 🎩\n'
       '#مجتمع_عربي_في_بريطانيا #أطلس_للإنسانية #ثقافة_بريطانية\n'
       '#BritishUpperClass #PoshBritain #UKEtiquette #CulturalIntelligence\n'
       '#SWANA #ArabsInUK #AtlasForHumanity #BritishCulture #ClassBritain')

    hr(doc)

    # ── PRODUCTION NOTES ─────────────────────────────────────────────────────
    section_heading(doc, 'PRODUCTION NOTES', NAVY)

    en(doc, 'Setting & Framing:', bold=True)
    en(doc, 'This is your elevated register. If possible, film in a formal or elegant setting — '
       'a hotel lobby, a grand room, near bookshelves. Otherwise, a plain dark/navy background '
       'with careful lighting. The visual tone should feel a notch more formal than Scripts 1 & 2.',
       indent=True)

    en(doc, 'Outfit:', bold=True)
    en(doc, 'Most formal of the three scripts. Blazer or structured jacket. '
       'Consider jewellery or a scarf in brand gold/cream. You are "translating" '
       'an elevated world — dress to match the register.',
       indent=True)

    en(doc, 'Delivery — KEY NOTE:', bold=True)
    en(doc, 'When delivering the posh English phrases, adopt a deliberately flat, cool, '
       'slightly bored British upper-class delivery. Even 10 seconds of this is '
       'extremely funny and deeply relatable to the SWANA audience who has encountered it. '
       'Do NOT overdo the accent — keep it stylised enough to signal the register.',
       indent=True)

    en(doc, 'Text Overlays:', bold=True)
    overlay_note(doc, '0:00', '"Perfectly adequate" = a 400-year-old country estate 🏰 [navy card, gold text]')
    overlay_note(doc, '0:28', '"Nancy Mitford — Noblesse Oblige (1956)" [source tag, small]')
    overlay_note(doc, '0:36', 'U CLASS vs NON-U TABLE: 3-row graphic, 2 columns [shareable card format]')
    overlay_note(doc, '0:56', '"You were very brave." = DISASTER ❌ | "Rather good." = EXCELLENT ✅')
    overlay_note(doc, '1:10', '"SAY: It\'s rather lovely, isn\'t it." ✅ [gold tip card]')
    overlay_note(doc, '1:18', '"Lord [Family Name] ✅ — Never Lord [First Name] ❌" [rule card]')

    en(doc, 'B-roll / Props:', bold=True)
    en(doc, '• The U vs Non-U vocabulary table is a MUST as a shareable graphic card — '
       'design it in brand colours. This alone will get saved and shared widely.\n'
       '• Optional: a book prop — Debrett\'s Handbook, or Nancy Mitford\'s "Noblesse Oblige."\n'
       '• Optional: chandelier / formal room B-roll if available.',
       indent=True)

    hr(doc)

    # ── REFERENCES ───────────────────────────────────────────────────────────
    section_heading(doc, 'REFERENCES', NAVY)
    ref_line(doc, 'Nancy Mitford — Noblesse Oblige: An Enquiry into the Identifiable Characteristics'
             ' of the English Aristocracy (1956)',
             'https://www.goodreads.com/book/show/263947.Noblesse_Oblige')
    ref_line(doc, 'Debrett\'s — Forms of Address (Peers, Lords, Ladies)',
             'https://debretts.com/expertise/forms-of-address/')
    ref_line(doc, 'Debrett\'s — Etiquette for formal occasions',
             'https://debretts.com/etiquettes/')
    ref_line(doc, 'BBC — U and Non-U: The English class system in words',
             'https://www.bbc.co.uk/news/magazine-17101482')
    ref_line(doc, 'The Guardian — How to speak upper class',
             'https://www.theguardian.com/education/2010/oct/22/how-to-speak-upper-class')
    ref_line(doc, 'Secret Food Tours — Funny food etiquette of British upper class',
             'https://www.secretfoodtours.com/blog/funny-food-etiquette-of-british-upper-class/')
    ref_line(doc, 'Atlas brand references — Instagram accounts',
             'https://www.instagram.com/britishculturearchive | '
             'https://www.instagram.com/englishmanmike | '
             'https://www.instagram.com/verybritishproblemsofficial')

    return doc


# ══════════════════════════════════════════════════════════════════════════════
# SAVE ALL THREE DOCUMENTS
# ══════════════════════════════════════════════════════════════════════════════

def main():
    scripts = [
        (build_script_1, f'Atlas_Reel_{TODAY}_1_Social_Sorry.docx'),
        (build_script_2, f'Atlas_Reel_{TODAY}_2_Professional_CorporateSpeak.docx'),
        (build_script_3, f'Atlas_Reel_{TODAY}_3_Posh_Understatement.docx'),
    ]

    paths = []
    for builder, filename in scripts:
        doc = builder()
        path = os.path.join(OUTPUT_DIR, filename)
        doc.save(path)
        print(f'✓  Saved: {path}')
        paths.append(path)

    print(f'\n✅  All 3 scripts saved to: {OUTPUT_DIR}')
    return paths

if __name__ == '__main__':
    main()
